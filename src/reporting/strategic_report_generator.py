"""
Strategic Report Generator v1.0
--------------------------------

Amaç:
- Strategic Control Panel Agent tarafından üretilen önerileri
  insan onayı sonrası Word (docx) + Excel (xlsx) formatına dönüştürmek.
"""

from docx import Document
from openpyxl import Workbook

from agentic.memory.long_term_memory import LongTermMemory
from governance.audit_logger import log_event


class StrategicReportGenerator:
    def __init__(self):
        self.mem = LongTermMemory()

    def load_proposal(self, proposal_id: str):
        for p in self.mem.memory.get("approval_queue", []):
            if p["id"] == proposal_id:
                return p
        return None

    def export_word(self, proposal, path="Strategic_Plan.docx"):
        doc = Document()
        doc.add_heading("AI4OHS-HYBRID Strategic Package", level=1)

        doc.add_heading("OHS Strategy", level=2)
        doc.add_paragraph(proposal["content"]["ohs_strategy"]["strategy_markdown"])

        doc.add_heading("Agent Network Evolution", level=2)
        doc.add_paragraph(
            str(proposal["content"]["agent_network_evolution"]["topology_suggestion"])
        )

        doc.add_heading("RAG vNext Proposal", level=2)
        doc.add_paragraph(str(proposal["content"]["rag_vNext"]["rag_vNext_proposal"]))

        doc.save(path)
        return path

    def export_excel(self, proposal, path="RAG_vNext.xlsx"):
        wb = Workbook()
        ws = wb.active
        ws.title = "RAG vNext"

        ws.append(["Key", "Value"])
        rag_next = proposal["content"]["rag_vNext"]["rag_vNext_proposal"]
        ws.append(["rag_vNext", rag_next])

        wb.save(path)
        return path

    def generate(self, proposal_id: str):
        proposal = self.load_proposal(proposal_id)

        if proposal is None:
            raise ValueError("Proposal ID not found.")

        if proposal["status"] != "APPROVED":
            raise PermissionError("Cannot generate report. Proposal is not approved.")

        word_path = self.export_word(proposal)
        excel_path = self.export_excel(proposal)

        log_event("REPORT_GEN", "Strategic report generated.", {"proposal_id": proposal_id})

        return {
            "word": word_path,
            "excel": excel_path,
            "message": "Reports generated successfully.",
        }
