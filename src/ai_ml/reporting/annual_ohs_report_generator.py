"""
AI4OHS-HYBRID
Annual OHS Report Generator

Bağlantı:
- TRIR / LTIFR / Severity Index → ohs_kpi_dashboard.compute_kpis + plot_kpi_dashboard
- ESS + 6331 uyum heatmap → compliance_heatmap.generate_heatmap_matrix

Üretir:
- Annual_OHS_Report.docx
- Annual_OHS_Report.xlsx
"""

import os
from typing import Dict, List

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import PatternFill

from ai_ml.reporting.ohs_kpi_dashboard import compute_kpis, plot_kpi_dashboard
from governance.compliance_heatmap import HEATMAP_COLORS, generate_heatmap_matrix


def generate_annual_excel(
    kpis: Dict[str, Dict], ess_items: List[Dict], filename: str = "Annual_OHS_Report.xlsx"
) -> str:
    wb = Workbook()

    # Sheet 1: KPI
    ws_kpi = wb.active
    ws_kpi.title = "KPI"

    headers = [
        "Period",
        "Hours",
        "Total Recordable",
        "LTI",
        "Lost Days",
        "TRIR",
        "LTIFR",
        "SeverityIndex",
    ]
    ws_kpi.append(headers)

    for period, vals in sorted(kpis.items()):
        ws_kpi.append(
            [
                period,
                vals["hours"],
                vals["total_recordable"],
                vals["lti"],
                vals["lost_days"],
                vals["TRIR"],
                vals["LTIFR"],
                vals["SeverityIndex"],
            ]
        )

    # Sheet 2: ESS–6331 Heatmap
    ws_heat = wb.create_sheet(title="ESS_6331_Heatmap")
    matrix = generate_heatmap_matrix(ess_items)

    all_laws = sorted({law for vals in matrix.values() for law in vals.keys()})
    ess_list = sorted(matrix.keys())

    ws_heat.cell(row=1, column=1, value="6331 / ESS")
    for j, ess in enumerate(ess_list, start=2):
        ws_heat.cell(row=1, column=j, value=ess)

    for i, law in enumerate(all_laws, start=2):
        ws_heat.cell(row=i, column=1, value=law)
        for j, ess in enumerate(ess_list, start=2):
            status = matrix.get(ess, {}).get(law, "")
            cell = ws_heat.cell(row=i, column=j, value=status)
            if status in HEATMAP_COLORS:
                cell.fill = PatternFill("solid", fgColor=HEATMAP_COLORS[status])

    os.makedirs(os.path.dirname(filename) or ".", exist_ok=True)
    wb.save(filename)
    return filename


def generate_annual_word(
    kpis: Dict[str, Dict],
    ess_items: List[Dict],
    kpi_png: str,
    filename: str = "Annual_OHS_Report.docx",
) -> str:
    doc = Document()
    doc.add_heading("Annual OHS Performance Report", level=1)

    # KPI summary
    doc.add_heading("1. KPI Overview (TRIR / LTIFR / Severity Index)", level=2)

    if os.path.exists(kpi_png):
        doc.add_paragraph("Figure 1 – OHS KPI trends (TRIR / LTIFR / Severity Index).")
        doc.add_picture(kpi_png, width=Inches(6))

    # KPI table
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Period"
    hdr[1].text = "TRIR"
    hdr[2].text = "LTIFR"
    hdr[3].text = "Severity Index"
    hdr[4].text = "Recordable Cases"

    for period, vals in sorted(kpis.items()):
        row = table.add_row().cells
        row[0].text = period
        row[1].text = f"{vals['TRIR']:.2f}"
        row[2].text = f"{vals['LTIFR']:.2f}"
        row[3].text = f"{vals['SeverityIndex']:.2f}"
        row[4].text = str(vals["total_recordable"])

    # ESS–6331 section
    doc.add_page_break()
    doc.add_heading("2. ESS and 6331 Compliance Heatmap", level=2)

    matrix = generate_heatmap_matrix(ess_items)
    ess_list = sorted(matrix.keys())
    all_laws = sorted({law for vals in matrix.values() for law in vals.keys()})

    htable = doc.add_table(rows=1, cols=len(ess_list) + 1)
    hh = htable.rows[0].cells
    hh[0].text = "Law / ESS"
    for i, ess in enumerate(ess_list, start=1):
        hh[i].text = ess

    for law in all_laws:
        row = htable.add_row().cells
        row[0].text = law
        for i, ess in enumerate(ess_list, start=1):
            status = matrix.get(ess, {}).get(law, "")
            row[i].text = status

    os.makedirs(os.path.dirname(filename) or ".", exist_ok=True)
    doc.save(filename)
    return filename


def generate_annual_ohs_report(
    incident_records: List[Dict],
    ess_items: List[Dict],
    period_hours: Dict[str, float] = None,
    out_dir: str = "reports/annual",
) -> Dict[str, str]:
    os.makedirs(out_dir, exist_ok=True)

    # KPI hesapla
    kpis = compute_kpis(incident_records, period_hours)
    kpi_png = os.path.join(out_dir, "ohs_kpi_dashboard.png")
    plot_kpi_dashboard(kpis, kpi_png)

    # Excel + Word üret
    excel_path = os.path.join(out_dir, "Annual_OHS_Report.xlsx")
    word_path = os.path.join(out_dir, "Annual_OHS_Report.docx")

    generate_annual_excel(kpis, ess_items, excel_path)
    generate_annual_word(kpis, ess_items, kpi_png, word_path)

    return {"excel": excel_path, "word": word_path, "kpi_png": kpi_png}
