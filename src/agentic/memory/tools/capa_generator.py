"""
AI4OHS-HYBRID – CAPA/NCR Auto Generator v2.0
Generates Excel or DOCX corrective action sheets.
"""

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import PatternFill


def severity_color(sev):
    if sev >= 15:
        return "FF0000"
    if sev >= 8:
        return "FFA500"
    return "FFFF00"


def generate_capa_excel(capa_items, filename="CAPA_NCR.xlsx"):
    wb = Workbook()
    ws = wb.active
    ws.title = "CAPA_NCR"

    headers = ["Issue", "Root Cause", "Corrective Action", "Responsible", "Due Date", "Severity"]
    ws.append(headers)

    for c in capa_items:
        row = [
            c["issue"],
            c.get("root_cause", ""),
            c["action"],
            c.get("responsible", ""),
            c.get("due", ""),
            c.get("severity", 1),
        ]
        ws.append(row)

        sev = c.get("severity", 1)
        ws.cell(ws.max_row, 6).fill = PatternFill("solid", fgColor=severity_color(sev))

    wb.save(filename)
    return filename


def generate_capa_docx(capa_items, filename="CAPA_NCR.docx"):
    doc = Document()
    doc.add_heading("CAPA / NCR Summary", level=1)

    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr_texts = ["Issue", "Root Cause", "Corrective Action", "Responsible", "Due Date", "Severity"]
    for i, h in enumerate(hdr_texts):
        hdr[i].text = h

    for c in capa_items:
        row = table.add_row().cells
        row[0].text = c["issue"]
        row[1].text = c.get("root_cause", "")
        row[2].text = c["action"]
        row[3].text = c.get("responsible", "")
        row[4].text = c.get("due", "")
        row[5].text = str(c.get("severity", 1))

    doc.save(filename)
    return filename
