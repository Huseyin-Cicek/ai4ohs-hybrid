"""
Report Writer – Template Automation v2.0
"""

from docx import Document

from agentic.memory.long_term_memory import LongTermMemory

memory = LongTermMemory()


def load_template(template_id: str):
    tpls = memory.memory.get("templates", {}).get("reports", [])
    for t in tpls:
        if t["template_id"] == template_id:
            return t
    return None


def build_report(template_id, findings, compliance_map, capa_list):
    tpl = load_template(template_id)
    doc = Document()

    doc.add_heading(tpl["name"], level=1)

    for section in tpl["sections"]:
        doc.add_heading(section, level=2)

        if section.lower() == "findings":
            for f in findings:
                doc.add_paragraph(f)

        if section.lower() == "compliance mapping":
            for c in compliance_map:
                doc.add_paragraph(f"{c['ref']}: {c['desc']}")

        if section.lower() == "capa":
            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Issue"
            hdr[1].text = "Corrective"
            hdr[2].text = "Timeline"
            for c in capa_list:
                row = table.add_row().cells
                row[0].text = c["issue"]
                row[1].text = c["action"]
                row[2].text = c["date"]

    return doc
    return doc
